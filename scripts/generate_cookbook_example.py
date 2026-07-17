#!/usr/bin/env python3
"""Generate the fictional DOCX source used by the README cookbook."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph

DEFAULT_OUTPUT = Path("examples/cookbook/aurora_ai_complaint_triage_process.docx")

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EDF4FA"
PALE_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
AMBER = "FFF2CC"


def _shade(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _cell_margins(
    cell: _Cell, *, top: int = 90, start: int = 90, bottom: int = 90, end: int = 90
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_header(row: _Row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _cant_split(row: _Row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_cell_text(cell: _Cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.2)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def _table(
    document: DocumentObject, headers: list[str], rows: list[list[str]], widths: list[float]
) -> Table:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    header = table.rows[0]
    _repeat_header(header)
    for index, value in enumerate(headers):
        _set_cell_text(header.cells[index], value, bold=True, color=WHITE)
        _shade(header.cells[index], NAVY)
        header.cells[index].width = Inches(widths[index])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        _cant_split(row)
        for column_index, value in enumerate(values):
            _set_cell_text(row.cells[column_index], value)
            row.cells[column_index].width = Inches(widths[column_index])
            if row_index % 2:
                _shade(row.cells[column_index], PALE_GRAY)
    return table


def _set_repeat_table_layout(table: Table) -> None:
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _add_field(paragraph: Paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def _configure_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string("222B45")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, color in (
        ("Title", 28, NAVY),
        ("Subtitle", 13, MID_GRAY),
        ("Heading 1", 17, NAVY),
        ("Heading 2", 12, BLUE),
        ("Heading 3", 10.5, NAVY),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.font.name = "Aptos"
        callout.font.size = Pt(9)
        callout.font.color.rgb = RGBColor.from_string(NAVY)
        callout.paragraph_format.left_indent = Inches(0.18)
        callout.paragraph_format.right_indent = Inches(0.18)
        callout.paragraph_format.space_before = Pt(4)
        callout.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = "AURORA FINANCIAL GROUP  |  FICTIONAL TRAINING DOCUMENT"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MID_GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Internal training example  |  Page ")
    run.font.name = "Aptos"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)
    _add_field(footer, "PAGE")
    footer.add_run(" of ")
    _add_field(footer, "NUMPAGES")


def _add_cover(document: DocumentObject) -> None:
    document.add_paragraph("OPERATING PROCESS", style="Subtitle")
    title = document.add_paragraph(
        "AI-Assisted Customer Complaint Triage and Escalation", style="Title"
    )
    title.paragraph_format.space_after = Pt(12)
    subtitle = document.add_paragraph(
        "A deliberately imperfect source document for the Document Enhancer cookbook",
        style="Subtitle",
    )
    subtitle.paragraph_format.space_after = Pt(20)

    banner = document.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_text(
        banner.cell(0, 0),
        "FICTIONAL - NO CUSTOMER DATA - SAFE FOR LOCAL OR LIVE DEMONSTRATION",
        bold=True,
        color=NAVY,
    )
    _shade(banner.cell(0, 0), LIGHT_BLUE)

    document.add_paragraph("Document control", style="Heading 2")
    _table(
        document,
        ["Document ID", "Version", "Status", "Owner", "Effective date", "Classification"],
        [
            [
                "PROC-AURORA-CCT-001",
                "0.8",
                "Draft for pilot",
                "Complaint Operations Manager",
                "2026-08-03",
                "Internal",
            ]
        ],
        [1.35, 0.65, 1.1, 1.55, 1.0, 1.0],
    )

    p = document.add_paragraph(style="Callout")
    p.add_run("Why this is a useful demonstration: ").bold = True
    p.add_run(
        "the document contains real Word headings, lists, captions, and native tables; explicit roles, "
        "controls, decisions, evidence, and metrics; and four intentional conflicts that should be routed "
        "to human review instead of silently resolved."
    )
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PALE_BLUE)
    p_pr.append(shading)

    document.add_paragraph(
        "Aurora Financial Group, Aurora Bank N.A., all systems, people, committees, thresholds, and "
        "records in this document are fictional. This source is designed to demonstrate structure "
        "recovery, analysis, clarification, governed rewrite, audit, semantic export, and grounded RAG."
    )


def _add_governance(document: DocumentObject) -> None:
    heading = document.add_heading("1. Document governance", level=1)
    heading.paragraph_format.page_break_before = True
    _table(
        document,
        ["Field", "Current draft value", "Notes"],
        [
            ["Business owner", "ROLE-COMPLAINT-OPS-MANAGER", "First-line process owner"],
            [
                "Accountable executive",
                "ROLE-HEAD-CUSTOMER-CARE",
                "Accountable for customer outcomes",
            ],
            ["Independent challenger", "ROLE-COMPLIANCE-COMPLAINTS", "Second-line challenge"],
            ["Approving authority", "COMMITTEE-CUSTOMER-RISK", "Pilot approval pending"],
            [
                "Legal entities",
                "Aurora Bank N.A.; Aurora Digital Services LLC",
                "United States only",
            ],
            [
                "Authoritative repository",
                "SYS-AURORA-GRC-001",
                "Signed approval and version history",
            ],
            [
                "Next review",
                "2027-08-02",
                "Earlier review after material model or regulatory change",
            ],
        ],
        [1.6, 2.5, 2.9],
    )

    document.add_heading("2. Purpose", level=1)
    document.add_paragraph(
        "This process receives customer complaints from approved channels, validates the record, uses "
        "a bounded classification model to propose product, issue, severity, and routing labels, requires "
        "human review before any customer-impacting decision, and escalates high-risk matters. The intended "
        "outcome is a complete, timely, reproducible case record with clear ownership and evidence."
    )

    document.add_heading("3. Scope and applicability", level=1)
    document.add_paragraph("In scope:")
    for item in (
        "Complaints submitted through secure web, authenticated mobile, recorded call-center, and approved mail-intake channels.",
        "US customers of Aurora Bank N.A. and Aurora Digital Services LLC.",
        "Deposit, card, payment, digital-access, and servicing complaints received after the effective date.",
        "English and Spanish complaints when an approved translation service is available.",
    ):
        document.add_paragraph(item, style="List Bullet")
    document.add_paragraph("Out of scope:")
    for item in (
        "Fraud claims governed by PROC-AURORA-FRAUD-002, legal demands, employee grievances, and anonymous social-media posts.",
        "Automated customer responses, eligibility decisions, compensation decisions, or complaint closure by the model.",
        "Non-US legal entities and complaints requiring an unavailable language service.",
    ):
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("4. Definitions", level=1)
    _table(
        document,
        ["Term", "Definition"],
        [
            [
                "Complaint",
                "An expression of dissatisfaction requiring a documented response or investigation.",
            ],
            [
                "Priority 1 (P1)",
                "Potential immediate customer harm, regulatory breach, systemic outage, or vulnerable-customer risk.",
            ],
            [
                "AI routing confidence",
                "The model confidence associated with the proposed product, issue, and destination queue.",
            ],
            [
                "Human confirmation",
                "A named specialist accepts or corrects all proposed labels before downstream assignment.",
            ],
            [
                "Material batch action",
                "Reassignment, suppression, or closure affecting more than 25 complaint records.",
            ],
            ["Business day", "A day on calendar CAL-AURORA-US-001 in America/New_York."],
        ],
        [1.8, 5.2],
    )


def _add_roles_inputs(document: DocumentObject) -> None:
    document.add_heading("5. Roles and responsibilities", level=1)
    _table(
        document,
        ["Role ID", "Line", "Responsibilities", "Decision rights", "Escalation"],
        [
            [
                "ROLE-INTAKE-SPECIALIST",
                "First",
                "Validate intake; confirm or correct AI labels; record evidence",
                "May assign one case; may not approve batch actions",
                "Complaint Operations Manager",
            ],
            [
                "ROLE-COMPLAINT-OPS-MANAGER",
                "First",
                "Own workflow, staffing, controls, and recovery",
                "Approves routine overrides and operational resumption",
                "Head of Customer Care",
            ],
            [
                "ROLE-COMPLIANCE-COMPLAINTS",
                "Second",
                "Challenge rules, review P1 and regulatory cases, monitor themes",
                "May require escalation or remediation",
                "Chief Compliance Officer",
            ],
            [
                "ROLE-MODEL-OWNER",
                "First",
                "Maintain model, thresholds, inventory, and performance evidence",
                "Recommends model changes; cannot self-approve",
                "Model Risk Committee",
            ],
            [
                "ROLE-MODEL-RISK",
                "Second",
                "Independent model validation and change challenge",
                "Approves material model change within mandate",
                "Model Risk Committee",
            ],
            [
                "ROLE-PRIVACY-OFFICER",
                "Second",
                "Review privacy incidents and sensitive-data handling",
                "Directs containment and notification assessment",
                "Chief Privacy Officer",
            ],
        ],
        [1.35, 0.55, 2.25, 1.65, 1.2],
    )

    document.add_heading("6. Preconditions, triggers, and inputs", level=1)
    document.add_paragraph(
        "The process begins when an approved channel creates a complaint record. Processing requires an "
        "available case-management system, an active routing taxonomy, a deployed approved model version, "
        "and an assigned intake specialist. The service clock starts at channel receipt, not at model completion."
    )
    _table(
        document,
        ["Input ID", "Source", "Freshness and quality rule", "Owner", "Fallback"],
        [
            [
                "DATA-COMPLAINT-001",
                "SYS-CASE-001",
                "Created at receipt; unique case ID; channel timestamp present",
                "Customer Care Data Steward",
                "Stop case and create data-quality incident",
            ],
            [
                "DATA-CUSTOMER-001",
                "SYS-CUSTOMER-001",
                "Current identity and contact preferences; no full payment credentials",
                "Customer Data Owner",
                "Proceed with minimum verified data; restrict access",
            ],
            [
                "TAXONOMY-CCT-001",
                "SYS-GRC-001",
                "Effective version and digest must match model release",
                "Complaint Operations Manager",
                "Use last approved version only under declared incident",
            ],
            [
                "MODEL-CCT-ROUTER-003",
                "SYS-MODEL-REGISTRY-001",
                "Approved version, health check passed, no emergency suspension",
                "Model Owner",
                "Manual triage queue",
            ],
        ],
        [1.3, 1.25, 2.45, 1.45, 1.65],
    )


def _add_process(document: DocumentObject) -> None:
    document.add_heading("7. Process overview", level=1)
    document.add_paragraph(
        "The model is advisory. A person validates the source, reviews the proposed labels and rationale, "
        "and accepts or corrects the route. P1 indicators bypass confidence-based routing and are sent to "
        "the priority queue for human confirmation."
    )
    flow = document.add_table(rows=1, cols=9)
    flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    flow.autofit = False
    labels = [
        "Receive",
        "->",
        "Validate",
        "->",
        "Propose",
        "->",
        "Confirm",
        "->",
        "Assign / escalate",
    ]
    for index, label in enumerate(labels):
        _set_cell_text(
            flow.cell(0, index), label, bold=index % 2 == 0, color=WHITE if index % 2 == 0 else NAVY
        )
        _shade(flow.cell(0, index), BLUE if index % 2 == 0 else WHITE)
        flow.cell(0, index).width = Inches(0.92 if index % 2 == 0 else 0.25)
    caption = document.add_paragraph(
        "Figure 1. High-level complaint triage flow. The step and rule tables are authoritative."
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MID_GRAY)

    document.add_heading("8. Atomic process steps", level=1)
    step_table = _table(
        document,
        [
            "Step ID",
            "Performer",
            "Action",
            "System",
            "Output / evidence",
            "Timing",
            "Failure route",
        ],
        [
            [
                "STEP-CCT-010",
                "Intake specialist",
                "Verify channel, case ID, timestamp, customer identity status, consent flags, and minimum narrative",
                "SYS-CASE-001",
                "Validated intake checklist and source digest",
                "Within 10 minutes of receipt",
                "FAIL-CCT-DATA",
            ],
            [
                "STEP-CCT-020",
                "Routing service",
                "Propose product, issue, severity, destination, rationale, and confidence; store model version",
                "MODEL-CCT-ROUTER-003",
                "Immutable proposal record; no customer action",
                "Within 60 seconds",
                "FAIL-CCT-MODEL",
            ],
            [
                "STEP-CCT-030",
                "Intake specialist",
                "Review the narrative and proposal; accept or correct every label; record reason when corrected",
                "SYS-CASE-001",
                "Human-confirmed labels with reviewer ID and timestamp",
                "Before assignment",
                "FAIL-CCT-REVIEW",
            ],
            [
                "STEP-CCT-040",
                "Intake specialist",
                "Apply decision rules and assign the complaint to the approved queue",
                "SYS-CASE-001",
                "Queue assignment and routing evidence",
                "Routine: 4 business hours; P1: see SLA note",
                "FAIL-CCT-QUEUE",
            ],
            [
                "STEP-CCT-050",
                "Queue owner",
                "Acknowledge ownership and begin investigation; notify required oversight roles",
                "SYS-CASE-001",
                "Named owner, acknowledgement, due date, notifications",
                "P1: within 60 minutes of receipt",
                "FAIL-CCT-SLA",
            ],
            [
                "STEP-CCT-060",
                "Complaint operations manager",
                "Review daily exceptions, unassigned cases, overrides, and repeat failures",
                "SYS-CONTROL-001",
                "Signed daily control record and issue links",
                "By 12:00 next business day",
                "FAIL-CCT-CONTROL",
            ],
        ],
        [0.9, 1.05, 2.35, 1.15, 1.65, 1.25, 0.95],
    )
    _set_repeat_table_layout(step_table)

    document.add_heading("9. Decision rules and thresholds", level=1)
    rule_table = _table(
        document,
        ["Rule ID", "Condition", "Threshold", "Outcome", "Authority / override", "Evidence"],
        [
            [
                "RULE-CCT-001",
                "P1 indicator present",
                "Any approved P1 indicator",
                "Priority queue plus Compliance notification; human confirms",
                "No downgrade by model; specialist may add P1",
                "Indicator, rationale, reviewer action",
            ],
            [
                "RULE-CCT-002",
                "No P1 indicator and routing confidence is high",
                ">= 0.80",
                "Show proposed route for human confirmation",
                "Specialist may correct with reason",
                "Confidence, labels, reviewer decision",
            ],
            [
                "RULE-CCT-003",
                "Routing confidence below threshold",
                "< 0.80",
                "Manual triage queue; no proposed destination is treated as approved",
                "Complaint Operations Manager may reassign one case",
                "Queue receipt and final human labels",
            ],
            [
                "RULE-CCT-004",
                "Material batch action",
                "> 25 records",
                "Pause and obtain approval before execution",
                "Draft says manager approval; approval partner not stated",
                "Approved change record and before/after population",
            ],
            [
                "RULE-CCT-005",
                "Sensitive-data exposure suspected",
                "Any credible indicator",
                "Restrict case, notify Privacy, preserve evidence",
                "Privacy Officer directs containment",
                "Incident and access log IDs",
            ],
        ],
        [0.95, 1.65, 1.0, 1.9, 1.8, 1.4],
    )
    _set_repeat_table_layout(rule_table)


def _add_controls(document: DocumentObject) -> None:
    document.add_heading("10. Controls, risks, and evidence", level=1)
    control_table = _table(
        document,
        [
            "Control ID",
            "Risk / objective",
            "Procedure and population",
            "Owner / reviewer",
            "Threshold",
            "Evidence / retention",
        ],
        [
            [
                "CTRL-CCT-001",
                "RISK-CCT-MISROUTE: prevent unreviewed routing",
                "100% of cases require a named human confirmation before queue assignment",
                "Complaint Ops / Compliance sample review",
                "Zero unconfirmed assignments",
                "Proposal and confirmation log; 7 years after case closure",
            ],
            [
                "CTRL-CCT-002",
                "RISK-CCT-SLA: identify delayed P1 handling",
                "Daily review of all P1 cases received in prior 24 hours",
                "Complaint Ops / Compliance",
                "Human acknowledgement within 30 minutes of receipt",
                "Daily SLA report and exceptions; 7 years",
            ],
            [
                "CTRL-CCT-003",
                "RISK-CCT-DRIFT: detect model deterioration",
                "Monthly performance review by product, language, channel, and vulnerable-customer flag",
                "Model Owner / Model Risk",
                "Accuracy >= 92%; override rate <= 12%; no segment below 85%",
                "Monitoring pack and challenge; model life plus 7 years",
            ],
            [
                "CTRL-CCT-004",
                "RISK-CCT-ACCESS: protect complaint data",
                "Quarterly recertification of all privileged and queue-level access",
                "System Owner / Information Security",
                "100% reviewed; zero orphaned privileged access",
                "Access review package; 7 years",
            ],
        ],
        [1.0, 1.45, 2.4, 1.5, 1.35, 1.55],
    )
    _set_repeat_table_layout(control_table)

    document.add_heading("11. Exceptions, failure paths, escalation, and recovery", level=1)
    _table(
        document,
        ["Failure / exception", "Safe action", "Escalation", "Recovery and closure"],
        [
            [
                "FAIL-CCT-DATA: required intake field missing or corrupt",
                "Do not call the model; hold in restricted data-quality queue",
                "Data Steward immediately; manager after 30 minutes",
                "Correct source, rerun validation, retain before/after evidence",
            ],
            [
                "FAIL-CCT-MODEL: unavailable, unhealthy, or version mismatch",
                "Switch to manual triage; do not reuse stale proposals",
                "Model Owner and Technology Incident Manager",
                "Health check, approved release digest, controlled resumption",
            ],
            [
                "FAIL-CCT-REVIEW: no qualified reviewer available",
                "Keep case unassigned; preserve service clock",
                "Complaint Operations Manager immediately",
                "Named substitute reviewer and reason",
            ],
            [
                "FAIL-CCT-SLA: P1 acknowledgement late",
                "Continue customer protection actions; open control exception",
                "Compliance and Head of Customer Care",
                "Root cause, customer-impact assessment, remediation evidence",
            ],
            [
                "EXC-CCT-001: approved taxonomy rollback",
                "Use prior approved taxonomy with manual review of all affected cases",
                "Model Risk and Compliance",
                "Expires after 2 business days; reconcile every affected case",
            ],
        ],
        [1.65, 2.25, 1.7, 2.35],
    )

    document.add_heading("12. Outputs and downstream consumers", level=1)
    document.add_paragraph(
        "The process produces a validated complaint record, human-confirmed classification, queue owner, "
        "service due date, model and taxonomy version, decision evidence, notifications, and exception links. "
        "Customer Care uses the record for investigation; Compliance uses it for oversight and reporting; "
        "Model Risk uses de-identified monitoring extracts; Internal Audit receives read-only evidence."
    )


def _add_dependencies_metrics(document: DocumentObject) -> None:
    document.add_heading("13. Systems and dependencies", level=1)
    _table(
        document,
        ["Dependency ID", "Purpose / owner", "Criticality", "Service and recovery", "Continuity"],
        [
            [
                "SYS-CASE-001",
                "Complaint case management / Technology Product Owner",
                "Critical",
                "99.9% monthly; RTO 2h; RPO 15m",
                "Read-only intake export plus manual priority log",
            ],
            [
                "MODEL-CCT-ROUTER-003",
                "Advisory classification / Model Owner",
                "High",
                "RTO 30m; no customer decision authority",
                "Manual triage queue",
            ],
            [
                "SYS-MODEL-REGISTRY-001",
                "Approved model identity and evidence / Model Risk",
                "High",
                "RTO 4h; immutable releases",
                "Block model use if identity cannot be verified",
            ],
            [
                "TPSP-TRANSLATE-001",
                "Approved translation / Vendor Manager",
                "Moderate",
                "Contract SLA 4h; restricted data terms",
                "Qualified human translation or hold",
            ],
        ],
        [1.55, 2.0, 0.8, 1.65, 1.9],
    )

    document.add_heading("14. Metrics and monitoring", level=1)
    _table(
        document,
        ["Metric ID", "Definition and formula", "Target / limit", "Owner / forum", "Breach action"],
        [
            [
                "METRIC-CCT-001",
                "P1 acknowledged on time / all P1 received; weekly",
                "Target 100%; limit < 98%",
                "Complaint Ops / Customer Risk Committee",
                "Immediate case review; issue after 2 breaches in 30 days",
            ],
            [
                "METRIC-CCT-002",
                "Human-confirmed assignments / all assignments; daily",
                "100%; zero tolerance",
                "Complaint Ops / Controls Forum",
                "Disable automated queue action and investigate",
            ],
            [
                "METRIC-CCT-003",
                "Correct model proposals / validated sample; monthly",
                ">= 92% overall; >= 85% each monitored segment",
                "Model Owner / Model Risk Committee",
                "Threshold review, remediation, or model suspension",
            ],
            [
                "METRIC-CCT-004",
                "Human corrections / reviewed proposals; monthly",
                "Watch > 12%; limit > 18%",
                "Model Owner / Model Risk Committee",
                "Root-cause analysis and change assessment",
            ],
        ],
        [1.15, 2.45, 1.55, 1.75, 1.8],
    )

    document.add_heading("15. Records retention", level=1)
    document.add_paragraph(
        "Complaint source, model proposal, human confirmation, routing, control, exception, approval, and "
        "customer-response records are retained in SYS-CASE-001 and SYS-GRC-001 for seven years after case "
        "closure. Legal holds suspend disposition. Access is role-based and retrieval is tested annually."
    )
    document.add_paragraph(
        "The pilot readiness checklist still says five years after calendar-year end. Records Management "
        "must confirm which period is authoritative before approval."
    )


def _add_open_points(document: DocumentObject) -> None:
    heading = document.add_heading("16. Open points for pilot approval", level=1)
    heading.paragraph_format.page_break_before = True
    p = document.add_paragraph(style="Callout")
    p.add_run("Reviewer note: ").bold = True
    p.add_run(
        "These conflicts are intentional. The enhancement workflow should preserve them as questions and "
        "require evidence-backed human answers instead of choosing a value on its own."
    )
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), AMBER)
    p_pr.append(shading)

    _table(
        document,
        ["Open point", "Conflicting draft statements", "Required decision owner"],
        [
            [
                "P1 acknowledgement SLA",
                "STEP-CCT-050 says 60 minutes; CTRL-CCT-002 says 30 minutes",
                "Head of Customer Care with Compliance concurrence",
            ],
            [
                "AI routing confidence",
                "RULE-CCT-002 uses 0.80; pilot readiness note uses 0.85",
                "Model Risk Committee",
            ],
            [
                "Material batch approval",
                "RULE-CCT-004 names manager approval but does not identify required independent approval",
                "Complaint Operations and Compliance",
            ],
            [
                "Record retention",
                "Section 15 says 7 years after closure; readiness checklist says 5 years after year end",
                "Records Management with Legal",
            ],
        ],
        [1.45, 3.7, 2.35],
    )

    document.add_heading("17. Pilot readiness notes", level=1)
    for item in (
        "Use 0.85 confidence before a proposed route may be shown as high confidence.",
        "Require 30-minute human acknowledgement for P1 complaints.",
        "Retain pilot complaint records for 5 years after calendar-year end.",
        "Do not enable material batch actions until approval roles are confirmed.",
        "Run a bilingual fairness and accuracy review before Spanish-language activation.",
    ):
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("18. Version history and approvals", level=1)
    _table(
        document,
        ["Version", "Date", "Change", "Author", "Challenge", "Decision"],
        [
            [
                "0.6",
                "2026-06-12",
                "Initial operating draft",
                "Complaint Operations",
                "Not started",
                "Working draft",
            ],
            [
                "0.7",
                "2026-07-01",
                "Added model, privacy, and recovery controls",
                "Complaint Operations",
                "Compliance comments open",
                "Working draft",
            ],
            [
                "0.8",
                "2026-07-16",
                "Added pilot metrics and explicit open points",
                "Complaint Operations",
                "Model Risk review pending",
                "Draft for pilot",
            ],
        ],
        [0.7, 0.95, 2.4, 1.35, 1.4, 1.15],
    )

    document.add_heading("Appendix A. Source inventory", level=1)
    document.add_paragraph(
        "Fictional related documents: POL-AURORA-CUSTOMER-001 Customer Treatment Policy; "
        "STD-AURORA-MODEL-002 Model Use Standard; PROC-AURORA-INCIDENT-001 Incident Management; "
        "SCHED-AURORA-COMPLAINT-007 Records Schedule; and CAL-AURORA-US-001 Business Calendar."
    )


def build_document(output: Path) -> None:
    document = Document()
    _configure_document(document)
    _add_cover(document)
    _add_governance(document)
    _add_roles_inputs(document)
    _add_process(document)
    _add_controls(document)
    _add_dependencies_metrics(document)
    _add_open_points(document)

    properties = document.core_properties
    properties.title = "AI-Assisted Customer Complaint Triage and Escalation"
    properties.subject = "Fictional Document Enhancer cookbook source"
    properties.author = "Document Enhancer contributors"
    properties.keywords = "fictional, complaint triage, AI governance, controls, cookbook"
    properties.comments = "Generated by scripts/generate_cookbook_example.py"

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build_document(args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

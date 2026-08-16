"""Focused tests for the file-backed v2 runner."""

from __future__ import annotations

import base64
import io
import json
import zipfile
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT

from document_enhancer.core import CoreRunner, RunStore
from document_enhancer.core.integrity import RecipeConfigurationMismatchError
from document_enhancer.core.layout import (
    AUDIT,
    DECISIONS_YAML,
    DRAFT_AUDIT,
    DRAFT_DOCUMENT,
    DRAFT_DOCUMENT_DOCX,
    DRAFT_TRANSFORMATION,
    DRAFT_VISUAL_EXTRACTIONS,
    FINAL_DOCX,
    FINAL_MARKDOWN,
    HTML_REPORT,
    ORIGINAL_DOCUMENT_PREFIX,
    QUESTIONS_MARKDOWN,
    REVIEW,
    REVIEW_INDEX_MARKDOWN,
    REWRITE_PLAN,
    RUN_RECORD,
    SEAL,
    SOURCE_MARKDOWN,
    SOURCE_METADATA,
)
from document_enhancer.core.models import (
    AuditReport,
    Decision,
    FlowEdge,
    Question,
    ReviewReport,
    RewritePlan,
    RewritePlanItem,
    RunRecord,
    Section,
)
from document_enhancer.core.recipes import load_recipe
from document_enhancer.core.review import merge_provider_review
from document_enhancer.core.rewrite import (
    apply_reviewer_decisions,
    apply_template_stubs,
    render_docx,
)
from document_enhancer.core.store import register_artifact
from document_enhancer.ingest.pipeline import DocumentIngestor

PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _approve_all(run_path: Path) -> None:
    path = run_path / DECISIONS_YAML
    text = path.read_text(encoding="utf-8")
    path.write_text(
        text.replace("approve_rewrite: false", "approve_rewrite: true")
        .replace('answer: ""', "answer: approved")
        .replace("disposition: defer", "disposition: accept"),
        encoding="utf-8",
    )


@pytest.mark.unit
def test_runner_pauses_for_questions_and_resumes_from_yaml(tmp_path: Path) -> None:
    source = tmp_path / "input.md"
    source.write_text(
        "# Intake\n\nThe owner submits the request.\n\n# Decision\n\nStatus: TBD\n",
        encoding="utf-8",
    )
    runner = CoreRunner(tmp_path / "runs")

    waiting = runner.start(source)

    assert waiting.status == "waiting"
    assert waiting.phase == "human_review"
    assert waiting.unresolved_question_ids == ["question-placeholder-001"]
    assert all(
        (tmp_path / "runs" / waiting.run_id / path).is_file()
        for path in (
            DRAFT_TRANSFORMATION,
            DRAFT_DOCUMENT,
            DRAFT_DOCUMENT_DOCX,
            DRAFT_AUDIT,
            DRAFT_VISUAL_EXTRACTIONS,
        )
    )
    assert (tmp_path / "runs" / waiting.run_id / REVIEW_INDEX_MARKDOWN).is_file()
    assert (tmp_path / "runs" / waiting.run_id / QUESTIONS_MARKDOWN).is_file()
    assert (tmp_path / "runs" / waiting.run_id / HTML_REPORT).is_file()
    decisions = tmp_path / "runs" / waiting.run_id / DECISIONS_YAML
    decisions.write_text(
        "approve_rewrite: true\n"
        'steering: ""\n'
        "waivers: []\n"
        "decisions:\n"
        "  - question_id: question-placeholder-001\n"
        "    answer: approved\n"
        "    disposition: accept\n",
        encoding="utf-8",
    )

    complete = runner.resume(waiting.run_id)

    assert complete.status == "succeeded"
    assert complete.phase == "verify"
    assert (tmp_path / "runs" / complete.run_id / RUN_RECORD).stat().st_size < 50_000
    final = (tmp_path / "runs" / complete.run_id / FINAL_MARKDOWN).read_text(encoding="utf-8")
    assert "Status: approved" in final
    audit = json.loads((tmp_path / "runs" / complete.run_id / AUDIT).read_text())
    assert audit["status"] == "pass"
    assert set(complete.artifacts) >= {
        "source.original",
        "review.report",
        "output.final_markdown",
        "output.final_docx",
        "output.semantic",
        "output.ontology",
        "output.graph",
        "audit.report",
        "audit.changes",
        "audit.source_to_target",
    }
    retry = runner.start(source)
    assert retry.run_id != complete.run_id
    assert (tmp_path / "runs" / complete.run_id / RUN_RECORD).is_file()


@pytest.mark.unit
def test_runner_accepts_a_canonical_generated_suggestion(tmp_path: Path) -> None:
    source = tmp_path / "input.md"
    source.write_text(
        "The owner receives the request, reviews the evidence, and records the outcome.\n",
        encoding="utf-8",
    )
    runner = CoreRunner(tmp_path / "runs")
    waiting = runner.start(source)
    run_path = tmp_path / "runs" / waiting.run_id
    template = (run_path / DECISIONS_YAML).read_text(encoding="utf-8")

    assert 'question: "What are the intended major sections for this document?"' in template
    assert "suggestion:" in template
    answered = template.replace("approve_rewrite: false", "approve_rewrite: true").replace(
        "disposition: defer", "disposition: accept_suggestion"
    )
    (run_path / DECISIONS_YAML).write_text(answered, encoding="utf-8")

    complete = runner.resume(waiting.run_id)

    assert complete.phase == "verify"
    decisions = json.loads((run_path / "json/06-decisions.json").read_text(encoding="utf-8"))
    assert decisions["decisions"][0]["disposition"] == "accept_suggestion"
    final = (run_path / FINAL_MARKDOWN).read_text(encoding="utf-8")
    assert "purpose, scope, responsibilities" in final


@pytest.mark.unit
def test_template_stubs_match_markdown_headings_not_body_prose() -> None:
    recipe = load_recipe(
        Path(__file__).parents[3] / "reference_packs/enterprise_core",
        document_type="process",
    )
    plan = RewritePlan(
        recipe_id=recipe.recipe_id,
        source_digest="a" * 64,
        items=[
            RewritePlanItem(
                section_id="missing-sec-proc-metadata",
                title="Document metadata and governance",
                missing_required=True,
                requirement_id="SEC-PROC-METADATA",
            ),
            RewritePlanItem(
                section_id="missing-sec-proc-metrics",
                title="Metrics, service levels, and monitoring",
                missing_required=True,
                requirement_id="SEC-PROC-METRICS",
            ),
        ],
    )
    decisions = [
        Decision(
            question_id="question-required-sec-proc-metadata",
            answer="Owner-approved governance metadata.",
        ),
        Decision(
            question_id="question-required-sec-proc-metrics",
            answer="Owner-approved metrics and monitoring.",
        ),
    ]

    final, changes = apply_template_stubs(
        "# Existing section\n\nThe process records governance decisions and metrics daily.\n",
        plan=plan,
        recipe=recipe,
        decisions=decisions,
        waived_requirement_ids=set(),
    )

    assert "## Document metadata and governance" in final
    assert "## Metrics, service levels, and monitoring" in final
    assert len(changes) == 2


@pytest.mark.unit
def test_runner_completes_clean_structured_document(tmp_path: Path) -> None:
    source = tmp_path / "input.md"
    source.write_text(
        "# Intake\n\nThe owner submits the request.\n\n"
        "# Approval\n\nThe manager reviews and approves the request.\n",
        encoding="utf-8",
    )

    runner = CoreRunner(tmp_path / "runs")
    result = runner.start(source)
    assert result.status == "waiting"
    assert "audit.seal" not in result.artifacts
    _approve_all(tmp_path / "runs" / result.run_id)
    result = runner.resume(result.run_id)

    assert result.status == "succeeded"
    assert result.unresolved_question_ids == []
    assert result.artifacts["source.original"].sha256 == result.source_digest
    assert (tmp_path / "runs" / result.run_id / FINAL_DOCX).stat().st_size > 0
    assert "audit.seal" in result.artifacts
    with pytest.raises(RuntimeError, match="sealed"):
        RunStore(tmp_path / "runs").write_text(result.run_id, FINAL_MARKDOWN, "tampered")


@pytest.mark.unit
def test_recovery_uses_registered_decisions_not_mutable_yaml(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "input.md"
    source.write_text("# Source\n\nStatus: TBD\n", encoding="utf-8")
    runner = CoreRunner(tmp_path / "runs")
    waiting = runner.start(source)
    run_path = tmp_path / "runs" / waiting.run_id
    _approve_all(run_path)

    def interrupted(*_: object, **__: object) -> RunRecord:
        raise RuntimeError("simulated promotion interruption")

    monkeypatch.setattr(runner, "_finish", interrupted)
    with pytest.raises(RuntimeError, match="simulated promotion interruption"):
        runner.resume(waiting.run_id)

    registered = runner.store.load_run(waiting.run_id)
    assert "review.decisions" in registered.artifacts
    canonical = json.loads((run_path / "json/06-decisions.json").read_text(encoding="utf-8"))
    assert canonical["decisions"][0]["answer"] == "approved"
    runner.store.save_run(registered.model_copy(update={"status": "running", "phase": "rewrite"}))
    decisions = run_path / DECISIONS_YAML
    decisions.write_text(
        decisions.read_text(encoding="utf-8").replace("answer: approved", "answer: compromised"),
        encoding="utf-8",
    )
    monkeypatch.undo()

    recovered = runner.resume(waiting.run_id)

    assert recovered.status == "succeeded"
    assert "audit.seal" in recovered.artifacts
    final = (run_path / FINAL_MARKDOWN).read_text(encoding="utf-8")
    assert "Status: approved" in final
    assert "compromised" not in final


@pytest.mark.unit
@pytest.mark.parametrize(
    "runner_kwargs",
    [
        {"document_type": "desktop_procedure"},
        {"recipe_pack": Path(__file__).parents[3] / "reference_packs" / "enterprise_core"},
    ],
    ids=["changed-configuration", "changed-recipe"],
)
def test_recovery_blocks_changed_recipe_or_configuration(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    runner_kwargs: dict[str, Any],
) -> None:
    source = tmp_path / "input.md"
    source.write_text("# Source\n\nStatus: TBD\n", encoding="utf-8")
    runner = CoreRunner(tmp_path / "runs")
    waiting = runner.start(source)
    run_path = tmp_path / "runs" / waiting.run_id
    _approve_all(run_path)

    def interrupted(*_: object, **__: object) -> RunRecord:
        raise RuntimeError("simulated promotion interruption")

    monkeypatch.setattr(runner, "_finish", interrupted)
    with pytest.raises(RuntimeError, match="simulated promotion interruption"):
        runner.resume(waiting.run_id)
    monkeypatch.undo()

    registered = runner.store.load_run(waiting.run_id)
    runner.store.save_run(registered.model_copy(update={"status": "failed", "phase": "verify"}))
    changed_runner = CoreRunner(tmp_path / "runs", **runner_kwargs)

    with pytest.raises(RecipeConfigurationMismatchError):
        changed_runner.resume(waiting.run_id)

    assert not (run_path / SEAL).exists()


@pytest.mark.unit
def test_run_record_keeps_the_selected_provider_mode(tmp_path: Path) -> None:
    source = tmp_path / "input.md"
    source.write_text("# Intake\n\nThe owner reviews the result.\n", encoding="utf-8")

    result = CoreRunner(tmp_path / "runs", execution_mode="live").start(source)

    assert result.execution_mode == "live"
    assert RunStore(tmp_path / "runs").load_run(result.run_id).execution_mode == "live"


@pytest.mark.unit
def test_run_store_rejects_artifact_path_escape(tmp_path: Path) -> None:
    store = RunStore(tmp_path / "runs")
    store.create_dir("run-1")

    with pytest.raises(ValueError, match="escapes"):
        store.write_text("run-1", "../outside.txt", "no")


@pytest.mark.unit
def test_resume_rehydrates_an_interrupted_analyze_phase(tmp_path: Path) -> None:
    source = tmp_path / "input.md"
    source.write_text("# Review\n\nThe owner reviews the result.\n", encoding="utf-8")
    raw = DocumentIngestor().parse(source)
    store = RunStore(tmp_path / "runs")
    store.create_dir("interrupted")
    record = RunRecord(
        run_id="interrupted",
        source_digest=raw.source_digest,
        source_name=source.name,
        status="running",
        phase="analyze",
    )
    store.save_run(record)
    original_artifact = store.write_bytes(
        "interrupted",
        f"{ORIGINAL_DOCUMENT_PREFIX}.md",
        source.read_bytes(),
        media_type=raw.media_type,
    )
    normalized_artifact = store.write_text(
        "interrupted", SOURCE_MARKDOWN, source.read_text(encoding="utf-8")
    )
    record = register_artifact(record, "source.original", original_artifact)
    record = register_artifact(record, "source.normalized", normalized_artifact)
    store.save_run(record)

    resumed = CoreRunner(tmp_path / "runs").resume("interrupted")

    assert resumed.status == "waiting"
    assert resumed.phase == "human_review"


@pytest.mark.unit
def test_resume_retries_a_failed_verification_with_current_rewrite_contract(
    tmp_path: Path,
) -> None:
    class RewriteStub:
        calls = 0

        def rewrite(self, **_: object) -> tuple[str, list[str]]:
            self.calls += 1
            if self.calls == 1:
                return "# Source\n\nTBD\n", ["introduced unresolved placeholder"]
            return "# Source\n\nThe owner reviews the result.\n", ["resolved placeholder"]

    source = tmp_path / "input.md"
    source.write_text("# Source\n\nThe owner reviews the result.\n", encoding="utf-8")
    rewrite_stub = RewriteStub()
    runner = CoreRunner(tmp_path / "runs", rewrite_provider=rewrite_stub)

    waiting = runner.start(source)
    _approve_all(tmp_path / "runs" / waiting.run_id)
    resumed = runner.resume(waiting.run_id)

    assert resumed.status == "succeeded"
    assert "audit.seal" in resumed.artifacts
    assert rewrite_stub.calls == 0


@pytest.mark.unit
def test_provider_enrichment_and_rewrite_are_recorded_as_optional_artifacts(tmp_path: Path) -> None:
    class ReviewStub:
        def review(self, **_: object) -> ReviewReport:
            return ReviewReport(summary="provider review")

    class RewriteStub:
        def rewrite(self, **_: object) -> tuple[str, list[str]]:
            return "# Source\n\nOwner approved.\n", ["rewrote the body"]

    source = tmp_path / "input.md"
    source.write_text("# Source\n\nThe owner reviews the result.\n", encoding="utf-8")

    runner = CoreRunner(
        tmp_path / "runs", review_provider=ReviewStub(), rewrite_provider=RewriteStub()
    )
    result = runner.start(source)
    _approve_all(tmp_path / "runs" / result.run_id)
    result = runner.resume(result.run_id)

    assert result.status == "succeeded"
    assert (
        "The owner reviews the result."
        in (tmp_path / "runs" / result.run_id / FINAL_MARKDOWN).read_text()
    )
    assert (tmp_path / "runs" / result.run_id / REWRITE_PLAN).is_file()


@pytest.mark.unit
def test_provider_questions_are_part_of_the_human_gate(tmp_path: Path) -> None:
    class ReviewStub:
        def review(self, **_: object) -> ReviewReport:
            return ReviewReport(
                summary="provider review",
                questions=[
                    Question(
                        question_id="provider-q-1",
                        prompt="Confirm the owner.",
                        reason="The provider found an unresolved business choice.",
                    )
                ],
            )

    source = tmp_path / "input.md"
    source.write_text("# Source\n\nThe owner reviews the result.\n", encoding="utf-8")

    result = CoreRunner(tmp_path / "runs", review_provider=ReviewStub()).start(source)

    assert result.status == "waiting"
    assert result.unresolved_question_ids == ["provider-q-1"]


@pytest.mark.unit
def test_offline_rewrite_applies_natural_language_open_point_decisions() -> None:
    source = (
        "P1: within 60 minutes of receipt\n"
        "STEP-CCT-050 says 60 minutes; CTRL-CCT-002 says 30 minutes\n"
        "Draft says manager approval; approval partner not stated\n"
        "RULE-CCT-004 names manager approval but does not identify required independent approval\n"
        "The pilot readiness checklist still says five years after calendar-year end. "
        "Records Management must confirm which period is authoritative before approval.\n"
        "Retain pilot complaint records for 5 years after calendar-year end.\n"
        "Section 15 says 7 years after closure; readiness checklist says 5 years after year end\n"
    )
    decision = Decision(
        question_id="question-open-points-001",
        answer=(
            "Use 30 minutes, retain records for seven years after case closure, and require an "
            "independent Compliance approver for every material batch action."
        ),
    )

    rewritten, changes = apply_reviewer_decisions(source, decisions=[decision])

    assert "60 minutes" not in rewritten
    assert "5 years" not in rewritten
    assert "approval partner not stated" not in rewritten
    assert "independent Compliance concurrence is required" in rewritten
    assert len(changes) >= 7


@pytest.mark.unit
def test_docx_renderer_preserves_markdown_structure_and_native_tables() -> None:
    markdown = (
        "# Governed process\n\n"
        "## Responsibilities\n\n"
        "The **owner** reviews *evidence* and records `control-id`.\n\n"
        "- Validate the request\n"
        "- Record the outcome\n\n"
        "| Step | Owner | Input | Action | Evidence | Timing |\n"
        "| --- | --- | --- | --- | --- | --- |\n"
        "| Intake | Analyst | Request | Validate | Checklist | 10 minutes |\n"
    )

    document = Document(io.BytesIO(render_docx(markdown)))

    paragraphs = {paragraph.text: paragraph for paragraph in document.paragraphs}
    title_style = paragraphs["Governed process"].style
    section_style = paragraphs["Responsibilities"].style
    list_style = paragraphs["Validate the request"].style
    assert title_style is not None and title_style.name == "Heading 1"
    assert section_style is not None and section_style.name == "Heading 2"
    assert list_style is not None and list_style.name.startswith("List Bullet")
    body = paragraphs["The owner reviews evidence and records control-id."]
    assert any(run.text == "owner" and run.bold for run in body.runs)
    assert any(run.text == "evidence" and run.italic for run in body.runs)
    assert document.sections[0].orientation == WD_ORIENT.LANDSCAPE
    assert len(document.tables) == 1
    table = document.tables[0]
    assert [cell.text for cell in table.rows[0].cells] == [
        "Step",
        "Owner",
        "Input",
        "Action",
        "Evidence",
        "Timing",
    ]
    assert all(cell.paragraphs[0].runs[0].bold for cell in table.rows[0].cells)
    assert table.rows[1].cells[0].text == "Intake"


@pytest.mark.unit
def test_runner_preserves_source_screenshot_as_referenced_appendix(tmp_path: Path) -> None:
    image = tmp_path / "submit.png"
    image.write_bytes(PNG_1X1)
    source = tmp_path / "process.md"
    source.write_text(
        "# Submission\n\nSelect Submit to complete the request.\n\n"
        "![Submission confirmation](submit.png)\n",
        encoding="utf-8",
    )

    runner = CoreRunner(tmp_path / "runs")
    result = runner.start(source)
    _approve_all(tmp_path / "runs" / result.run_id)
    result = runner.resume(result.run_id)
    run_path = tmp_path / "runs" / result.run_id
    source_metadata = json.loads((run_path / SOURCE_METADATA).read_text(encoding="utf-8"))
    final_markdown = (run_path / FINAL_MARKDOWN).read_text(encoding="utf-8")
    audit = json.loads((run_path / AUDIT).read_text(encoding="utf-8"))

    assert result.status == "succeeded"
    assert [figure["figure_id"] for figure in source_metadata["figures"]] == ["FIG-001"]
    assert "Select Submit to complete the request. **[FIG-001]**" in final_markdown
    assert "## Appendix A — Source screenshots" in final_markdown
    assert "### [FIG-001] Submission confirmation" in final_markdown
    assert "![Submission confirmation](../assets/final/FIG-001.png)" in final_markdown
    assert (run_path / "assets/source/FIG-001.png").read_bytes() == PNG_1X1
    assert (run_path / "assets/final/FIG-001.png").read_bytes() == PNG_1X1
    with zipfile.ZipFile(run_path / FINAL_DOCX) as archive:
        assert any(name.startswith("word/media/") for name in archive.namelist())
    assert audit["checks"]["figure_references_valid"] is True
    assert audit["checks"]["figure_appendix_complete"] is True
    assert audit["checks"]["figure_asset_digests_match"] is True
    assert audit["checks"]["final_docx_figures_embedded"] is True
    assert "Source screenshots" in (run_path / HTML_REPORT).read_text(encoding="utf-8")


@pytest.mark.unit
def test_runner_carries_docx_screenshot_into_final_docx_appendix(tmp_path: Path) -> None:
    image = tmp_path / "review.png"
    image.write_bytes(PNG_1X1)
    source = tmp_path / "process.docx"
    document = Document()
    document.add_heading("Review", level=1)
    document.add_paragraph("Open the review screen and confirm the result.")
    document.add_picture(str(image))
    caption = document.add_paragraph("Review confirmation screen")
    caption.style = "Caption"
    document.save(str(source))

    runner = CoreRunner(tmp_path / "runs")
    result = runner.start(source)
    _approve_all(tmp_path / "runs" / result.run_id)
    result = runner.resume(result.run_id)
    run_path = tmp_path / "runs" / result.run_id
    final_markdown = (run_path / FINAL_MARKDOWN).read_text(encoding="utf-8")

    assert result.status == "succeeded"
    assert "confirm the result. **[FIG-001]**" in final_markdown
    assert "### [FIG-001] Review confirmation screen" in final_markdown
    with zipfile.ZipFile(run_path / FINAL_DOCX) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        assert media


@pytest.mark.unit
def test_review_provider_uses_one_macro_call_and_bounded_section_batches(tmp_path: Path) -> None:
    class ReviewStub:
        def __init__(self) -> None:
            self.batch_sizes: list[int] = []

        def review(self, **_: object) -> ReviewReport:
            self.batch_sizes.append(0)
            return ReviewReport(summary="macro")

        def review_sections(self, *, sections: list[object], **_: object) -> ReviewReport:
            self.batch_sizes.append(len(sections))
            return ReviewReport(summary="sections")

    source = tmp_path / "input.md"
    source.write_text(
        "\n\n".join(f"# Section {index}\n\nThe owner reviews step {index}." for index in range(9)),
        encoding="utf-8",
    )
    provider = ReviewStub()

    result = CoreRunner(tmp_path / "runs", review_provider=provider).start(source)

    assert result.status == "waiting"
    assert provider.batch_sizes == [0, 4, 4, 1]


@pytest.mark.unit
def test_flow_graph_requires_evidence_for_relationships(tmp_path: Path) -> None:
    source = tmp_path / "input.md"
    source.write_text(
        "# Intake\n\nThe owner records the request.\n\n"
        "# Approval\n\nIf complete, then submit to the manager.\n\n"
        "# Archive\n\nStore the approved record.\n",
        encoding="utf-8",
    )

    result = CoreRunner(tmp_path / "runs").start(source)
    review = ReviewReport.model_validate(
        json.loads((tmp_path / "runs" / result.run_id / REVIEW).read_text(encoding="utf-8"))
    )

    assert [node.node_id for node in review.flow_nodes] == [
        "section-001",
        "section-002",
        "section-003",
    ]
    assert {(edge.source, edge.target, edge.relation) for edge in review.flow_edges} == {
        ("section-002", "section-003", "branch"),
    }
    assert review.flow_edges[0].evidence_span_ids
    assert "section-001 --> section-002" not in review.mermaid


@pytest.mark.unit
def test_provider_flow_edges_are_evidence_filtered_and_typed() -> None:
    base = ReviewReport(
        summary="base",
        process_applicable=True,
        sections=[
            Section(section_id="one", title="One", level=1, span_ids=["span-1"]),
            Section(section_id="two", title="Two", level=1, span_ids=["span-2"]),
        ],
    )
    candidate = ReviewReport(
        summary="provider",
        flow_edges=[
            FlowEdge(
                edge_id="candidate-valid",
                source="one",
                target="two",
                relation="branch",
                evidence_span_ids=["span-1"],
            ),
            FlowEdge(
                edge_id="candidate-invalid",
                source="one",
                target="two",
                relation="sequence",
                evidence_span_ids=["unknown"],
            ),
        ],
    )

    merged = merge_provider_review(base, candidate, allowed_span_ids={"span-1", "span-2"})

    assert [(edge.source, edge.target, edge.relation) for edge in merged.flow_edges] == [
        ("one", "two", "branch")
    ]
    assert "branch" in merged.mermaid


@pytest.mark.unit
def test_provider_question_ids_remain_unique_across_section_batches() -> None:
    base = ReviewReport(
        summary="base",
        questions=[Question(question_id="q-001", prompt="Base?", reason="base")],
    )
    first = ReviewReport(
        summary="first",
        questions=[Question(question_id="q-001", prompt="First?", reason="first")],
    )
    second = ReviewReport(
        summary="second",
        questions=[Question(question_id="q-001", prompt="Second?", reason="second")],
    )

    merged = merge_provider_review(base, first, allowed_span_ids=set())
    merged = merge_provider_review(merged, second, allowed_span_ids=set())

    assert [item.question_id for item in merged.questions] == [
        "q-001",
        "llm-q-001",
        "llm-q-001-2",
    ]


@pytest.mark.unit
def test_independent_audit_provider_is_recorded_without_changing_offline_contract(
    tmp_path: Path,
) -> None:
    class AuditStub:
        def audit(self, **_: object) -> AuditReport:
            return AuditReport(
                status="pass",
                checks={"content_fidelity": True},
                summary="independent pass",
            )

    source = tmp_path / "input.md"
    source.write_text("# Source\n\nThe owner reviews the result.\n", encoding="utf-8")

    runner = CoreRunner(tmp_path / "runs", audit_provider=AuditStub())
    result = runner.start(source)
    _approve_all(tmp_path / "runs" / result.run_id)
    result = runner.resume(result.run_id)

    assert result.status == "succeeded"
    audit = json.loads((tmp_path / "runs" / result.run_id / AUDIT).read_text(encoding="utf-8"))
    assert audit["checks"]["independent_content_audit"] is True

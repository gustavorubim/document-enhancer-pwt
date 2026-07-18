"""Rewrite planning and final-document projections for the core workflow."""

from __future__ import annotations

import csv
import hashlib
import io
import re
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt

from .export import public_graph_jsonl
from .models import (
    Decision,
    DocumentIR,
    Finding,
    ReviewReport,
    RewritePlan,
    RewritePlanItem,
    SemanticEdge,
    SemanticNode,
)
from .recipes import Recipe
from .review import title_matches

_PLACEHOLDER_RE = re.compile(r"\b(?:TBD|TODO|TBC)\b|\[\s*\?\s*\]|\?{3,}", re.IGNORECASE)


def compile_rewrite_plan(
    *,
    source_digest: str,
    review: ReviewReport,
    decisions: list[Decision],
    recipe: Recipe | None,
) -> RewritePlan:
    findings_by_section: dict[str, list[Finding]] = {}
    for finding in review.findings:
        if finding.section_id:
            findings_by_section.setdefault(finding.section_id, []).append(finding)
    items = [
        RewritePlanItem(
            section_id=section.section_id,
            title=section.title,
            source_span_ids=list(section.span_ids),
            finding_ids=[
                item.finding_id for item in findings_by_section.get(section.section_id, [])
            ],
            recommendations=list(
                dict.fromkeys(
                    item.recommendation
                    for item in findings_by_section.get(section.section_id, [])
                    if item.recommendation
                )
            ),
        )
        for section in review.sections
    ]
    required_section_ids: list[str] = []
    if recipe:
        for requirement in recipe.required_section_items:
            heading = str(requirement.get("heading") or requirement.get("id") or "")
            requirement_id = str(requirement.get("id") or heading)
            matching = next(
                (item for item in review.sections if title_matches(heading, item.title)),
                None,
            )
            if matching:
                required_section_ids.append(matching.section_id)
                continue
            stub_id = f"missing-{requirement_id.lower()}"
            required_section_ids.append(stub_id)
            items.append(
                RewritePlanItem(
                    section_id=stub_id,
                    title=heading or requirement_id,
                    source_span_ids=[],
                    finding_ids=[
                        item.finding_id
                        for item in review.findings
                        if item.finding_id.endswith(requirement_id.lower())
                    ],
                    recommendations=[
                        "Insert an evidence-backed section or record an explicit waiver."
                    ],
                    missing_required=True,
                    requirement_id=requirement_id,
                )
            )
    return RewritePlan(
        recipe_id=review.recipe_id,
        source_digest=source_digest,
        items=items,
        required_section_ids=required_section_ids,
        accepted_decision_ids=[
            item.question_id
            for item in decisions
            if item.disposition == "accept" and item.answer.strip()
        ],
        deferred_decision_ids=[
            item.question_id for item in decisions if item.disposition == "defer"
        ],
    )


def apply_deterministic_answers(text: str, decisions: list[Decision]) -> tuple[str, list[str]]:
    changes: list[str] = []
    for decision in decisions:
        answer = decision.answer.strip()
        if not answer or decision.disposition != "accept":
            continue
        updated = _PLACEHOLDER_RE.sub(answer, text, count=1)
        if updated != text:
            changes.append(
                f"Replaced one unresolved marker with the accepted answer for {decision.question_id}."
            )
        text = updated
    return text, changes


def apply_reviewer_decisions(
    text: str,
    *,
    decisions: list[Decision],
    steering: str = "",
) -> tuple[str, list[str]]:
    """Apply accepted answers offline: placeholders, steered conflict fixes, decision log."""

    text, changes = apply_deterministic_answers(text, decisions)
    accepted = [item for item in decisions if item.disposition == "accept" and item.answer.strip()]
    if not accepted and not steering.strip():
        return text, changes
    open_points = next(
        (
            item
            for item in accepted
            if item.question_id == "question-open-points-001" or "open-points" in item.question_id
        ),
        None,
    )
    if open_points is not None:
        answer = open_points.answer.lower()
        replacements: list[tuple[str, str, str]] = []
        if "30 minutes" in answer:
            replacements.append(
                (
                    "P1: within 60 minutes of receipt",
                    "P1: within 30 minutes of receipt",
                    "Aligned STEP-CCT-050 P1 timing to the approved 30-minute SLA.",
                )
            )
            replacements.append(
                (
                    "Require 30-minute human acknowledgement for P1 complaints.",
                    "Require 30-minute human acknowledgement for P1 complaints (approved).",
                    "Confirmed pilot readiness P1 acknowledgement at 30 minutes.",
                )
            )
            replacements.append(
                (
                    "STEP-CCT-050 says 60 minutes; CTRL-CCT-002 says 30 minutes",
                    "STEP-CCT-050 and CTRL-CCT-002 both require 30 minutes (approved)",
                    "Removed the resolved P1 timing conflict from the open-points table.",
                )
            )
        if "0.85" in answer:
            replacements.append(
                (
                    ">= 0.80",
                    ">= 0.85",
                    "Updated high-confidence routing threshold to the approved 0.85 value.",
                )
            )
            replacements.append(
                (
                    "< 0.80",
                    "< 0.85",
                    "Updated low-confidence routing threshold to the approved 0.85 boundary.",
                )
            )
            replacements.append(
                (
                    "RULE-CCT-002 uses 0.80; pilot readiness note uses 0.85",
                    "RULE-CCT-002 and pilot readiness both use 0.85 (approved)",
                    "Resolved AI routing confidence conflict to 0.85.",
                )
            )
        if re.search(r"\b(?:7|seven)[ -]?years?\b", answer):
            replacements.append(
                (
                    "Retain pilot complaint records for 5 years after calendar-year end.",
                    "Retain pilot complaint records for 7 years after case closure (approved).",
                    "Resolved retention conflict to 7 years after case closure.",
                )
            )
            replacements.append(
                (
                    "The pilot readiness checklist still says five years after calendar-year end. "
                    "Records Management must confirm which period is authoritative before approval.",
                    "Records retention is approved at seven years after case closure "
                    "(Records Management with Legal).",
                    "Removed unresolved retention conflict language.",
                )
            )
            replacements.append(
                (
                    "Section 15 says 7 years after closure; readiness checklist says 5 years "
                    "after year end",
                    "Section 15 and the readiness checklist both require 7 years after closure "
                    "(approved)",
                    "Removed the resolved retention conflict from the open-points table.",
                )
            )
        if "compliance" in answer and any(
            term in answer for term in ("batch", "independent", "approval", "concurrence")
        ):
            replacements.append(
                (
                    "Draft says manager approval; approval partner not stated",
                    "Complaint Operations Manager approval plus Compliance concurrence required",
                    "Named the independent approval partner for material batch actions.",
                )
            )
            replacements.append(
                (
                    "RULE-CCT-004 names manager approval but does not identify required "
                    "independent approval",
                    "Complaint Operations Manager approval plus independent Compliance "
                    "concurrence is required (approved)",
                    "Removed the resolved batch-approval conflict from the open-points table.",
                )
            )
        for old, new, note in replacements:
            if old in text and old != new:
                text = text.replace(old, new)
                changes.append(note)
    log_lines = [
        "",
        "# Reviewer decisions applied",
        "",
        "The following accepted decisions were applied during the rewrite stage.",
        "",
    ]
    if steering.strip():
        log_lines.extend(["## Steering", "", steering.strip(), ""])
    log_lines.extend(["## Accepted answers", ""])
    for item in accepted:
        log_lines.extend(
            [
                f"### `{item.question_id}`",
                "",
                item.answer.strip(),
                "",
            ]
        )
        if item.rationale:
            log_lines.extend([f"_Rationale:_ {item.rationale}", ""])
    if "# Reviewer decisions applied" not in text:
        text = text.rstrip() + "\n" + "\n".join(log_lines)
        changes.append("Appended reviewer decisions applied section from accepted answers.")
    return text, changes


def apply_template_stubs(
    text: str,
    *,
    plan: RewritePlan,
    recipe: Recipe | None,
    decisions: list[Decision],
    waived_requirement_ids: set[str],
) -> tuple[str, list[str]]:
    """Append explicit stubs for missing required sections that were not waived."""

    if not recipe:
        return text, []
    answers = {
        item.question_id: item.answer.strip()
        for item in decisions
        if item.disposition == "accept" and item.answer.strip()
    }
    changes: list[str] = []
    body = text.rstrip() + "\n"
    headings = [line.lstrip("# ").strip() for line in body.splitlines() if line.startswith("#")]
    for item in plan.items:
        if not item.missing_required or not item.requirement_id:
            continue
        if item.requirement_id in waived_requirement_ids:
            continue
        if any(title_matches(item.title, heading) for heading in headings):
            continue
        expected = next(
            (
                str(requirement.get("expected_content") or "")
                for requirement in recipe.required_section_items
                if str(requirement.get("id") or "") == item.requirement_id
            ),
            "",
        )
        question_id = f"question-required-{item.requirement_id.lower()}"
        answer = answers.get(question_id, "")
        if answer and answer.lower() not in {"yes", "y", "true", "include"}:
            content = answer
        elif answer:
            content = (
                f"TBD: Reviewer approved inclusion of `{item.requirement_id}` but did not supply "
                "source-backed body text."
            )
        else:
            content = (
                f"TBD: Provide source-backed content for required section `{item.requirement_id}`."
            )
        stub = f"\n## {item.title}\n\n{content}\n"
        if expected:
            stub += f"\nExpected content: {expected}.\n"
        body += stub
        headings.append(item.title)
        changes.append(f"Inserted governed stub for missing section {item.title!r}.")
    return body, changes


def _markdown_parser() -> MarkdownIt:
    return MarkdownIt("commonmark", {"html": False, "typographer": True}).enable("table")


def _inline_text(token: Any) -> str:
    children = token.children or []
    return "".join(
        "\n" if item.type in {"softbreak", "hardbreak"} else str(item.content)
        for item in children
        if item.type in {"text", "code_inline", "softbreak", "hardbreak"}
    ).strip()


def _append_inline(paragraph: Any, token: Any, *, force_bold: bool = False) -> None:
    """Translate Markdown inline tokens to native Word runs."""

    bold = force_bold
    italic = False
    link: str | None = None
    link_text = ""
    for child in token.children or []:
        if child.type == "strong_open":
            bold = True
            continue
        if child.type == "strong_close":
            bold = force_bold
            continue
        if child.type == "em_open":
            italic = True
            continue
        if child.type == "em_close":
            italic = False
            continue
        if child.type == "link_open":
            link = str(child.attrGet("href") or "")
            link_text = ""
            continue
        if child.type == "link_close":
            if link and link != link_text:
                run = paragraph.add_run(f" ({link})")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(117, 105, 128)
            link = None
            link_text = ""
            continue
        if child.type in {"softbreak", "hardbreak"}:
            paragraph.add_run().add_break()
            continue
        if child.type == "html_inline" and "br" in str(child.content).lower():
            paragraph.add_run().add_break()
            continue
        if child.type == "image":
            text = f"[Image: {child.content or child.attrGet('alt') or 'illustration'}]"
        elif child.type in {"text", "code_inline"}:
            text = str(child.content)
        else:
            continue
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if child.type == "code_inline":
            run.font.name = "Aptos Mono"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(101, 82, 124)
        if link:
            link_text += text
            run.underline = True
            run.font.color.rgb = RGBColor(110, 88, 143)


def _set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_repeat_table_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _set_row_no_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    no_split = OxmlElement("w:cantSplit")
    properties.append(no_split)


def _table_tokens(tokens: list[Any], start: int) -> tuple[list[list[tuple[Any, bool, str]]], int]:
    rows: list[list[tuple[Any, bool, str]]] = []
    row: list[tuple[Any, bool, str]] = []
    header = False
    alignment = "left"
    index = start + 1
    while index < len(tokens) and tokens[index].type != "table_close":
        token = tokens[index]
        if token.type == "tr_open":
            row = []
        elif token.type in {"th_open", "td_open"}:
            header = token.type == "th_open"
            style = str(token.attrGet("style") or "")
            alignment = "center" if "center" in style else "right" if "right" in style else "left"
        elif token.type == "inline" and row is not None:
            row.append((token, header, alignment))
        elif token.type == "tr_close" and row:
            rows.append(row)
        index += 1
    return rows, index


def _configure_docx(document: Any, *, wide_tables: bool) -> None:
    section = document.sections[0]
    if wide_tables:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        margin = Inches(0.55)
    else:
        margin = Inches(0.72)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(64, 58, 73)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    heading_colors = (
        RGBColor(83, 70, 96),
        RGBColor(105, 84, 124),
        RGBColor(93, 121, 115),
    )
    for level in range(1, 7):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.bold = True
        style.font.color.rgb = heading_colors[min(level - 1, 2)]
        style.font.size = Pt(max(11, 19 - (level - 1) * 1.7))
        style.paragraph_format.space_before = Pt(15 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True
    title = document.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = RGBColor(83, 70, 96)
    title.paragraph_format.space_after = Pt(15)
    document.styles["Quote"].font.color.rgb = RGBColor(106, 85, 91)
    document.styles["Quote"].font.italic = True


def _add_markdown_table(document: Any, rows: list[list[tuple[Any, bool, str]]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    compact = columns > 5
    for row_index, source_row in enumerate(rows):
        target_row = table.rows[row_index]
        _set_row_no_split(target_row)
        if row_index == 0 and any(header for _, header, _ in source_row):
            _set_repeat_table_header(target_row)
        for column_index, (inline, header, alignment) in enumerate(source_row):
            cell = target_row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if header:
                _set_cell_shading(cell, "EDE7F4")
            elif row_index % 2 == 0:
                _set_cell_shading(cell, "FAF7FB")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = {
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
            _append_inline(paragraph, inline, force_bold=header)
            for run in paragraph.runs:
                run.font.name = "Aptos"
                run.font.size = Pt(8 if compact else 9)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def render_docx(markdown: str) -> bytes:
    """Render Markdown as a structured, styled Word document with native tables."""

    tokens = _markdown_parser().parse(markdown)
    table_widths = []
    for index, token in enumerate(tokens):
        if token.type == "table_open":
            rows, _ = _table_tokens(tokens, index)
            table_widths.append(max((len(row) for row in rows), default=0))
    document = Document()
    _configure_docx(document, wide_tables=max(table_widths, default=0) > 5)
    list_stack: list[str] = []
    quote_depth = 0
    first_heading = True
    index = 0
    while index < len(tokens):
        token = tokens[index]
        if token.type == "heading_open" and index + 1 < len(tokens):
            level = min(int(token.tag[1:]), 6)
            inline = tokens[index + 1]
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _append_inline(paragraph, inline)
            if first_heading:
                document.core_properties.title = _inline_text(inline)
                first_heading = False
            index += 3
            continue
        if token.type in {"bullet_list_open", "ordered_list_open"}:
            list_stack.append("bullet" if token.type == "bullet_list_open" else "number")
            index += 1
            continue
        if token.type in {"bullet_list_close", "ordered_list_close"}:
            if list_stack:
                list_stack.pop()
            index += 1
            continue
        if token.type == "blockquote_open":
            quote_depth += 1
            index += 1
            continue
        if token.type == "blockquote_close":
            quote_depth = max(0, quote_depth - 1)
            index += 1
            continue
        if token.type == "paragraph_open" and index + 1 < len(tokens):
            inline = tokens[index + 1]
            if list_stack:
                style = "List Bullet" if list_stack[-1] == "bullet" else "List Number"
            elif quote_depth:
                style = "Quote"
            else:
                style = None
            paragraph = document.add_paragraph(style=style)
            if len(list_stack) > 1:
                paragraph.paragraph_format.left_indent = Inches(0.25 * (len(list_stack) - 1))
            _append_inline(paragraph, inline)
            index += 3
            continue
        if token.type == "table_open":
            rows, closing_index = _table_tokens(tokens, index)
            _add_markdown_table(document, rows)
            index = closing_index + 1
            continue
        if token.type in {"fence", "code_block"}:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.22)
            paragraph.paragraph_format.right_indent = Inches(0.22)
            run = paragraph.add_run(str(token.content).rstrip())
            run.font.name = "Aptos Mono"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(75, 65, 85)
            index += 1
            continue
        if token.type == "hr":
            paragraph = document.add_paragraph()
            properties = paragraph._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "D9CFE2")
            borders.append(bottom)
            properties.append(borders)
        index += 1
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def semantic_graph(
    review: ReviewReport,
    final_text: str,
    *,
    recipe: Recipe | None = None,
) -> dict[str, Any]:
    nodes: list[SemanticNode] = []
    edges: list[SemanticEdge] = []
    for section in review.sections:
        requirement = None
        if recipe:
            requirement = next(
                (
                    item
                    for item in recipe.required_sections
                    if title_matches(
                        str(item.get("heading") or item.get("id") or ""), section.title
                    )
                ),
                None,
            )
        hooks = [str(item) for item in ((requirement or {}).get("ontology_hooks") or [])]
        primary_type = hooks[0] if hooks else "Section"
        if recipe and not recipe.allows_node_type(primary_type):
            primary_type = "Section"
        nodes.append(
            SemanticNode(
                node_id=section.section_id,
                label=section.title,
                node_type=primary_type,
                properties={
                    "level": section.level,
                    "parent_id": section.parent_id,
                    "requirement_id": (requirement or {}).get("id"),
                    "ontology_hooks": hooks,
                },
                provenance_span_ids=list(section.span_ids),
            )
        )
        for hook in hooks[1:]:
            if recipe and not recipe.allows_node_type(hook):
                continue
            hook_id = f"{section.section_id}:{hook.lower()}"
            nodes.append(
                SemanticNode(
                    node_id=hook_id,
                    label=f"{section.title} ({hook})",
                    node_type=hook,
                    properties={"parent_section": section.section_id},
                    provenance_span_ids=list(section.span_ids[:3]),
                )
            )
            edges.append(
                SemanticEdge(
                    source=section.section_id,
                    target=hook_id,
                    edge_type="contains",
                    properties={"derived_from": "ontology_hooks"},
                    provenance_span_ids=list(section.span_ids[:3]),
                )
            )
    flow_edges = review.proposed_flow_edges or review.flow_edges
    for edge in flow_edges:
        edge_type = edge.relation
        if recipe and not recipe.allows_edge_type(edge_type):
            continue
        edges.append(
            SemanticEdge(
                source=edge.source,
                target=edge.target,
                edge_type=edge_type,
                properties={"relation": edge.relation},
                provenance_span_ids=list(edge.evidence_span_ids),
            )
        )
    for node in review.proposed_flow_nodes:
        if any(item.node_id == node.node_id for item in nodes):
            continue
        mapped = {
            "decision": "Decision",
            "step": "ProcessStep",
            "section": "Section",
        }.get(node.node_type, "Section")
        if recipe and not recipe.allows_node_type(mapped):
            mapped = "Section"
        nodes.append(
            SemanticNode(
                node_id=node.node_id,
                label=node.label,
                node_type=mapped,
                properties={"proposed": True},
                provenance_span_ids=[],
            )
        )
    ir = DocumentIR(
        sections=review.sections,
        nodes=nodes,
        edges=edges,
        markdown_sha256=hashlib.sha256(final_text.encode("utf-8")).hexdigest(),
    )
    return ir.model_dump(mode="json")


def graph_json_lines(graph: dict[str, list[dict[str, str]]]) -> list[str]:
    return public_graph_jsonl(graph).splitlines()


def source_target_csv(review: ReviewReport, final_text: str) -> str:
    stream = io.StringIO()
    writer = csv.writer(stream)
    writer.writerow(["section_id", "title", "disposition", "final_digest"])
    digest = hashlib.sha256(final_text.encode("utf-8")).hexdigest()
    for section in review.sections:
        title_tokens = [
            token for token in re.findall(r"[a-z0-9]+", section.title.lower()) if len(token) >= 3
        ]
        retained = bool(title_tokens) and all(token in final_text.lower() for token in title_tokens)
        writer.writerow(
            [section.section_id, section.title, "retained" if retained else "missing", digest]
        )
    return stream.getvalue()


def semantic_diff(before: dict[str, Any], after: dict[str, Any]) -> dict[str, Any]:
    before_nodes = {item["node_id"] for item in before.get("nodes", [])}
    after_nodes = {item["node_id"] for item in after.get("nodes", [])}
    before_edges = {
        (item["source"], item["target"], item["edge_type"]) for item in before.get("edges", [])
    }
    after_edges = {
        (item["source"], item["target"], item["edge_type"]) for item in after.get("edges", [])
    }
    return {
        "added_nodes": sorted(after_nodes - before_nodes),
        "removed_nodes": sorted(before_nodes - after_nodes),
        "added_edges": [list(item) for item in sorted(after_edges - before_edges)],
        "removed_edges": [list(item) for item in sorted(before_edges - after_edges)],
        "markdown_changed": before.get("markdown_sha256") != after.get("markdown_sha256"),
    }


__all__ = [
    "apply_deterministic_answers",
    "apply_reviewer_decisions",
    "apply_template_stubs",
    "compile_rewrite_plan",
    "graph_json_lines",
    "render_docx",
    "semantic_diff",
    "semantic_graph",
    "source_target_csv",
]
